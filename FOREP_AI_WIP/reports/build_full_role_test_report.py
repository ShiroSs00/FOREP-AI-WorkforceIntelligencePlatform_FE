import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
RAW = sorted(REPORTS.glob("full_role_test_*.json"), key=lambda p: p.stat().st_mtime)[-1]
ROUTE_SMOKE = REPORTS / "frontend_route_smoke_latest.json"
OUT = REPORTS / "FOREP_Full_Role_Test_Report.docx"


def load_json(path, fallback=None):
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8-sig"))
    return fallback


raw = load_json(RAW, {})
routes = load_json(ROUTE_SMOKE, [])


def status_label(item):
    if not item:
        return "NOT TESTED"
    if item.get("ok") is True:
        return "PASS"
    if item.get("status") in (401, 403):
        return "ACCESS"
    return "FAIL"


def count_tests(role_data):
    tests = role_data.get("tests") or {}
    passed = sum(1 for v in tests.values() if v.get("ok") is True)
    failed = sum(1 for v in tests.values() if v.get("ok") is not True)
    return passed, failed, len(tests)


def clean_message(item):
    if not item:
        return ""
    msg = item.get("message") or item.get("error") or ""
    details = item.get("details")
    if isinstance(details, str) and details and details not in msg:
        msg = f"{msg} | {details}"
    return str(msg).replace("\n", " ").strip()


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "F2F4F7")
        set_cell_text(cell, text, bold=True, color="0B2545")
        if widths:
            cell.width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text)
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(31, 77, 120)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.10

for style_name, size, color in [
    ("Heading 1", 16, "2E74B5"),
    ("Heading 2", 13, "2E74B5"),
    ("Heading 3", 12, "1F4D78"),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("FOREP Full Role Functional Test Report")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(11, 37, 69)

subtitle = doc.add_paragraph()
subtitle.add_run("AI Workforce Intelligence Platform - frontend + backend API role verification").italic = True

tested_at = raw.get("testedAt", "")
base_url = raw.get("baseUrl") or "https://forep-ai-workforceintelligenceplatform.onrender.com"
generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

add_table(
    doc,
    ["Item", "Value"],
    [
        ["Backend base URL", base_url],
        ["Raw role-test JSON", str(RAW)],
        ["Frontend route smoke JSON", str(ROUTE_SMOKE)],
        ["API test timestamp", tested_at],
        ["Document generated", generated_at],
    ],
    [2.0, 4.5],
)

doc.add_heading("1. Executive Summary", level=1)
roles = raw.get("roles", {})
summary_rows = []
for role_name in ["ADMIN", "MANAGER", "HR", "EMPLOYEE"]:
    role = roles.get(role_name, {})
    passed, failed, total = count_tests(role)
    expected_actual = role.get("actualRole") or "Unknown"
    auth_ok = role.get("register", {}).get("ok") and role.get("login", {}).get("ok")
    role_note = "Role mapping OK"
    if role_name == "HR" and expected_actual != "HR":
        role_note = "Backend returned EMPLOYEE for HR registration/login context"
    if failed:
        role_note = f"{role_note}; {failed} API issue(s) remain"
    summary_rows.append([
        role_name,
        "PASS" if auth_ok else "FAIL",
        expected_actual,
        f"{passed}/{total} endpoint checks passed",
        role_note,
    ])

add_table(
    doc,
    ["Role", "Auth", "Actual backend role", "API result", "Note"],
    summary_rows,
    [0.85, 0.75, 1.2, 1.4, 2.3],
)

add_callout(
    doc,
    "Summary",
    "Frontend build and lint passed. Local route smoke passed after starting the Vite dev server. "
    "Admin, Manager and Employee role APIs are mostly usable. The main blockers are backend-side: HR is mapped as EMPLOYEE, "
    "AI generation returns 500, and project task reads fail after webhook/import creation.",
)

doc.add_heading("2. Frontend Route Smoke", level=1)
route_rows = []
for item in routes:
    route_rows.append([
        item.get("route", ""),
        "PASS" if item.get("ok") else "FAIL",
        str(item.get("status") or ""),
        item.get("message") or item.get("error") or "",
    ])
if not route_rows:
    route_rows.append(["Local frontend", "NOT TESTED", "", "No latest route smoke file was found."])
add_table(doc, ["Route", "Result", "HTTP", "Message"], route_rows, [1.3, 0.8, 0.7, 3.7])

doc.add_heading("3. Role Account and API Matrix", level=1)
for role_name in ["ADMIN", "MANAGER", "HR", "EMPLOYEE"]:
    role = roles.get(role_name, {})
    doc.add_heading(role_name.title(), level=2)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Test email", role.get("email", "")],
            ["Register", f"{status_label(role.get('register'))} ({role.get('register', {}).get('status', '')})"],
            ["Login", f"{status_label(role.get('login'))} ({role.get('login', {}).get('status', '')})"],
            ["Auth /me", f"{status_label(role.get('me'))} ({role.get('me', {}).get('status', '')})"],
            ["Actual role returned", role.get("actualRole", "Unknown")],
            ["Employee ID", role.get("employeeId", "Not available")],
        ],
        [2.0, 4.5],
    )

    rows = []
    for name, result in (role.get("tests") or {}).items():
        rows.append([
            name,
            status_label(result),
            str(result.get("status") or ""),
            "" if result.get("count") is None else str(result.get("count")),
            clean_message(result),
        ])
    add_table(doc, ["Capability / endpoint group", "Result", "HTTP", "Count", "Message"], rows, [2.15, 0.75, 0.55, 0.55, 2.5])

doc.add_heading("4. Backend Issues Found During Testing", level=1)
issues = [
    [
        "B1",
        "High",
        "HR role registration/login maps to EMPLOYEE.",
        "The HR account registered successfully, login worked, but /auth/me returned actualRole=EMPLOYEE.",
        "Check RegisterRequest role handling and user role persistence for HR / People Ops.",
    ],
    [
        "B2",
        "High",
        "Employee list and leave list fail for HR account.",
        "HR calls to /api/v1/employees and /api/v1/leaves returned 500 in this run.",
        "Fix authorization/role mapping first, then retest HR people operations endpoints.",
    ],
    [
        "B3",
        "High",
        "Project task read fails after GitHub/Jira webhook/import.",
        "Admin test passed project creation and webhooks, but tasksByProjectAfterWebhook returned 500.",
        "Check task serialization/imported task entity mapping after webhook processing.",
    ],
    [
        "B4",
        "Medium",
        "AI generate endpoint returns 500 for Employee account.",
        "POST /api/v1/ai/generate/{employeeId} failed while my-insights and AI runtime/suggestions worked.",
        "Inspect AI service request path, employee context, Ollama/FastAPI exception and backend error body.",
    ],
    [
        "B5",
        "Medium",
        "Unsupported global GET endpoints should not be used by FE.",
        "Backend returns method-not-supported for global GET /projects and GET /integrations; FE has been adjusted to use scoped endpoints.",
        "Keep Swagger explicit about scoped project/integration endpoints or add supported list endpoints later.",
    ],
]
add_table(doc, ["ID", "Severity", "Finding", "Evidence", "Suggested fix"], issues, [0.45, 0.75, 1.65, 2.0, 1.65])

doc.add_heading("5. Role-by-Role Retest Checklist", level=1)
checklist_rows = [
    ["Admin", "Login as admin, open Dashboard, Organizations, Teams, Projects, Integrations, AI Insights.", "Dashboard counts load from /admin/dashboard; CRUD uses real API; no unsupported GET list calls."],
    ["Manager", "Login as manager, open Dashboard, Team Tasks, Teams, Attendance, Leave, Analytics, AI Insights.", "Managed-team endpoints return data or clean empty states; no Admin-only controls."],
    ["HR", "Create/login HR account after backend role fix, open Employees, Attendance, Leave, Recruitment, People Analytics.", "Actual backend role is HR; employees/leaves load without 500."],
    ["Employee", "Login employee, open Dashboard, My Tasks, Attendance, Leave, Profile, AI Insights.", "Only personal endpoints are called; check-in/out and leave creation call real API."],
    ["Integrations", "Use Team ID, create GitHub/Jira config, run sync/logs/webhook test.", "No global GET /integrations; task import can be read from scoped project endpoint after backend fix."],
    ["AI", "Check runtime, suggestions, managed/my insights, generate insight.", "AI generate should return 200 or a domain-specific validation error, not raw 500."],
]
add_table(doc, ["Area", "What to test", "Pass condition"], checklist_rows, [1.0, 2.65, 2.85])

doc.add_heading("6. Frontend Verification Notes", level=1)
notes = [
    "npm.cmd run lint: passed.",
    "npm.cmd run build: passed.",
    "Local route smoke after restarting Vite: all tested SPA routes returned HTTP 200.",
    "The FE should continue hiding raw endpoint labels from user-facing cards where possible; technical details belong inside ErrorState details only.",
    "The FE must keep using scoped project/integration endpoints because the backend does not support global list GET endpoints for those controllers.",
]
for note in notes:
    p = doc.add_paragraph(style=None)
    p.style = styles["Normal"]
    p.paragraph_format.left_indent = Inches(0.25)
    p.add_run(note)

doc.add_heading("7. Raw Evidence", level=1)
add_callout(
    doc,
    "Evidence files",
    f"Raw API evidence is stored at {RAW}. Frontend route smoke evidence is stored at {ROUTE_SMOKE}. "
    "Use these files when comparing future backend redeploys.",
    fill="E8EEF5",
)

doc.save(OUT)
print(OUT)
