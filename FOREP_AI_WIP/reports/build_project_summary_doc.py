from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "FOREP_Project_Current_Summary.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(9.5)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade_cell(cell, "F2F4F7")
        set_cell(cell, h, bold=True, color="0B2545")
        if widths:
            cell.width = Inches(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return t


def callout(doc, title, body, fill="E8EEF5"):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(31, 77, 120)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.10
for style_name, size, color in [
    ("Heading 1", 16, "2E74B5"),
    ("Heading 2", 13, "2E74B5"),
    ("Heading 3", 12, "1F4D78"),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("FOREP Project Current Summary")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(11, 37, 69)

subtitle = doc.add_paragraph()
subtitle.add_run("AI Workforce Intelligence Platform - frontend handoff and status brief").italic = True

table(doc, ["Item", "Current value"], [
    ["Generated at", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
    ["Frontend stack", "React 19, Vite 8, Tailwind CSS 4, React Router 7, anime.js, lucide-react"],
    ["Backend base URL", "https://forep-ai-workforceintelligenceplatform.onrender.com"],
    ["API mode", "VITE_DATA_MODE=api; no mock fallback in normal app mode"],
    ["Local storage usage", "JWT token, selected role preference, theme preference, language preference only"],
    ["Latest verification", "npm.cmd run lint passed; npm.cmd run build passed; SPA route smoke returned HTTP 200"],
], [2.0, 4.5])

doc.add_heading("1. Product Positioning", level=1)
callout(
    doc,
    "FOREP in one line",
    "FOREP is an AI Workforce Intelligence Platform. It connects tasks, attendance, leave, projects, integrations, workload analytics and AI insights into role-based operational intelligence. The positioning remains: not a surveillance tool, but a transparency tool.",
)

doc.add_heading("2. Frontend Architecture", level=1)
table(doc, ["Area", "Implementation"], [
    ["Routing", "Public routes: /, /login, /register. Protected app routes render inside AppLayout with Sidebar, TopHeader and main Outlet."],
    ["Layouts", "PublicLayout is used for public pages. AppLayout keeps Sidebar and TopHeader stable so only main content changes."],
    ["Auth", "authService calls /api/v1/auth/login, /register, /logout and /me. JWT is stored in forep_auth_token and attached by apiClient."],
    ["API client", "apiClient uses VITE_API_BASE_URL, 15s timeout, Authorization header, JSON parsing, clean ApiError objects and backend message extraction."],
    ["Response mapping", "responseNormalizer handles direct arrays/objects, data/result/content/items/payload and safe field helpers."],
    ["Theme", "ThemeContext supports light/dark/system using class-based dark mode and forep_theme."],
    ["Language", "LanguageContext supports English/Vietnamese using forep_language."],
    ["Role state", "RoleContext derives selected UI role from /auth/me after login. Admin can preview roles; non-admin accounts should follow backend role."],
], [1.35, 5.15])

doc.add_heading("3. Routes and Main Pages", level=1)
table(doc, ["Route", "Page", "Purpose"], [
    ["/", "LandingPage", "Cinematic marketing landing page with scrollytelling SaaS positioning."],
    ["/login", "LoginPage", "Real backend login, no demo login."],
    ["/register", "RegisterPage", "Real backend registration."],
    ["/dashboard", "DashboardPage", "Role-based dashboard shell."],
    ["/tasks", "TaskPage", "Scoped task APIs: my, managed, project, team, status, sprint, reporter, organization, employee."],
    ["/projects", "ProjectsPage", "Project create/update and organization/team-scoped project reads."],
    ["/teams", "TeamPage", "Team CRUD, members, assign employee, join request, membership approve/end active."],
    ["/sprints", "SprintPage", "Sprint CRUD and active/organization scoped sprint reads."],
    ["/integrations", "IntegrationsPage", "Provider config create/update/delete, connect, sync, logs, runtime status and OAuth links."],
    ["/analytics", "AnalyticsPage", "Workload history by my, managed, organization, team or employee scope."],
    ["/ai-insights", "AIInsightsPage", "AI runtime, insights, suggestions, adopt suggestion and generate insight."],
    ["/attendance", "AttendancePage", "Check-in/out and role-scoped attendance history."],
    ["/leave", "LeavePage", "Leave create, approve/reject and role-scoped leave history."],
    ["/notifications", "NotificationPage", "Notification list/read/delete/read-all through backend API."],
    ["/settings", "SettingsPage", "Theme, language and profile/config placeholders."],
], [1.1, 1.55, 3.85])

doc.add_heading("4. Role-Based UI Scope", level=1)
table(doc, ["Role", "Primary modules", "Notes"], [
    ["Admin", "Dashboard, Organizations, Users, Teams, Projects, Employees, Tasks, Sprints, Events, Analytics, AI Insights, Integrations, Monitoring, Attendance, Leave, Notifications", "Platform/system control role. Admin can preview other roles for testing."],
    ["Manager", "Dashboard, Organization, Teams, Projects, Team Tasks, Sprints, Employees, Team Analytics, AI Insights, Events, Attendance, Leave, Integrations, Reports, Notifications", "Team operations role. Uses managed-team endpoints where available."],
    ["HR / People Ops", "Dashboard, Employees, Teams, Projects, Task Overview, Attendance, Leave, Recruitment, People Analytics, AI Insights, Reports, Notifications", "People operations role. Backend currently maps HR account as EMPLOYEE in test, so this needs backend fix."],
    ["Employee", "Dashboard, My Tasks, Projects, Attendance, Leave, Personal Analytics, Event Timeline, AI Insights, Notifications, Profile", "Personal workspace role. Uses my-* and profile endpoints."],
], [1.0, 3.25, 2.25])

doc.add_heading("5. Service Layer Coverage", level=1)
table(doc, ["Service", "Mapped capability"], [
    ["authService", "Login/register/logout/me/OAuth links and backend OAuth redirect URL helpers."],
    ["dashboardService", "Admin dashboard and employee dashboard endpoints."],
    ["organizationService", "Organization list/create/get/update/delete."],
    ["employeeService", "Employee list/get/update/delete/profile/team/org scoped reads."],
    ["teamService", "Team CRUD, assign employee, members, organization/manager/managed scopes, join request, membership approve/end-active."],
    ["projectService", "Project create/get/update and organization/team scoped reads. Delete is explicitly unavailable because Swagger has no DELETE project endpoint."],
    ["sprintService", "Sprint CRUD, active sprint and organization-scoped sprint APIs."],
    ["taskService", "Task CRUD, status update, assess, and all scoped task reads."],
    ["taskCommentService", "Task comment list/create/delete."],
    ["attendanceService", "Check-in/out and my/team/org/managed/employee attendance history."],
    ["leaveService", "Leave create, approve/reject and my/team/status/org/managed/employee leave reads."],
    ["notificationService", "List, unread, unread count, mark read, read all, delete."],
    ["analyticsService", "Workload history my/managed/team/org/employee and development mock generator endpoint."],
    ["aiInsightService", "Runtime status, generate insight, my/managed/team/org/employee insights."],
    ["aiSuggestionService", "All/scoped suggestions and adopt suggestion."],
    ["integrationService", "Integration config create/update/delete, connect, team reads, runtime status, sync, logs, GitHub/Jira webhook helpers."],
], [1.55, 4.95])

doc.add_heading("6. Latest API Test Snapshot", level=1)
table(doc, ["Metric", "Result"], [
    ["Swagger operations found", "119"],
    ["Full API probe calls", "103"],
    ["Passed", "87"],
    ["Failed", "16"],
    ["Latest probe files", "reports/full_api_probe_20260621000244.txt and .json"],
    ["Frontend route smoke", "Routes /, /login, /register, /dashboard, /tasks, /projects, /teams, /sprints, /integrations, /analytics, /ai-insights, /attendance, /leave, /notifications, /settings returned HTTP 200."],
], [2.0, 4.5])

doc.add_heading("7. Known Backend/API Issues", level=1)
table(doc, ["Severity", "Issue", "Current impact"], [
    ["High", "HR account maps to EMPLOYEE after registration/login.", "HR sidebar/dashboard/API scope cannot be fully validated until backend role persistence is fixed."],
    ["High", "GET /api/v1/tasks and GET /api/v1/tasks/managed-teams returned 500 in full API probe.", "FE avoids global task list and provides scoped selectors. Manager task API may still show ErrorState if backend returns 500."],
    ["High", "Leave list/status/approve/reject returned 500 in probe.", "LeavePage will show backend error messages for affected operations."],
    ["Medium", "OAuth redirects for GitHub/Google/Jira returned 503.", "OAuth link UI exists, but backend redirect is not ready."],
    ["Medium", "AI generate insight returned 503.", "AIInsightsPage calls real endpoint and shows backend error instead of faking generated content."],
    ["Medium", "Integration config create for GitHub/Jira returned 500 in probe with test payload.", "Integrations UI is wired, but backend validation/config handling needs review."],
    ["Medium", "DELETE sprint returned 500 for test sprint.", "Sprint delete UI calls real API; backend may reject depending on relations/state."],
], [0.75, 2.6, 3.15])

doc.add_heading("8. Deployment and Testing Notes", level=1)
table(doc, ["Topic", "Note"], [
    ["Environment", "Set VITE_API_BASE_URL=https://forep-ai-workforceintelligenceplatform.onrender.com and VITE_DATA_MODE=api."],
    ["Vercel SPA reload", "vercel.json should rewrite routes to /index.html so refreshing protected routes does not return 404."],
    ["Build", "npm.cmd run build passes. Vite warns that JS chunk is above 500 KB; this is a warning, not a build failure."],
    ["Manual test order", "Login/register, admin dashboard, organizations, teams, projects, sprints, integrations, tasks by scope, attendance, leave, analytics, AI insights, notifications."],
    ["Safe testing", "Use dedicated test accounts/resources. Avoid deleting shared organizations with teams because backend may enforce FK constraints."],
], [1.5, 5.0])

doc.add_heading("9. Recommended Next Steps", level=1)
for item in [
    "Fix backend HR role mapping first; it blocks accurate HR role UI validation.",
    "Fix backend 500 errors for global/managed task endpoints or keep FE scoped-only until stable.",
    "Fix leave approval/rejection and leave list endpoints for HR/manager use cases.",
    "Stabilize integration config create/connect for GitHub and Jira with real required credentials.",
    "Stabilize AI generate endpoint and return domain-specific validation messages instead of 503/500.",
    "After backend fixes, rerun reports/full_api_probe.ps1 and update the handoff report.",
]:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.25)
    p.add_run(item)

doc.save(OUT)
print(OUT)
