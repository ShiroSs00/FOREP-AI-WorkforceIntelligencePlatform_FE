from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("reports/FOREP_System_Overview_FE_BE_Action_Plan.docx")


BLUE = "0EA5E9"
DARK = "0F172A"
MUTED = "64748B"
LIGHT = "F8FAFC"
BORDER = "D7DEE8"
WARN = "FEF3C7"
RISK = "FEE2E2"
GOOD = "DCFCE7"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=DARK, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_fill(hdr[i], "EAF6FE")
        set_cell_text(hdr[i], header, bold=True, color="075985", size=8.5)
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8.5)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_callout(doc, title, body, fill="EAF6FE"):
    p = doc.add_paragraph()
    shade_paragraph(p, fill)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)
    p.add_run(body)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 6),
        ("Heading 2", 13, "0369A1", 10, 4),
        ("Heading 3", 11.5, "075985", 8, 3),
    ]:
        s = styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)


def main():
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(2)
    tr = title.add_run("FOREP System Overview & FE/BE Action Plan")
    tr.bold = True
    tr.font.size = Pt(23)
    tr.font.color.rgb = RGBColor.from_string(DARK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    sr = subtitle.add_run("AI Workforce Intelligence Platform - frontend/backend status, API changes, and fixes needed")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor.from_string(MUTED)

    meta = [
        ["Prepared for", "FOREP project"],
        ["Prepared on", datetime.now().strftime("%Y-%m-%d %H:%M")],
        ["Frontend commit reviewed", "78fea7d"],
        ["Frontend stack", "React 19, Vite 8, Tailwind CSS 4, React Router 7, anime.js"],
        ["Backend base URL", "https://forep-ai-workforceintelligenceplatform.onrender.com"],
        ["Primary backend source", "Swagger + FRONTEND_BACKEND_CHANGELOG.md"],
    ]
    add_table(doc, ["Field", "Value"], meta, widths=[1.8, 4.7])

    doc.add_heading("1. Executive Summary", level=1)
    add_callout(
        doc,
        "Current status",
        "FOREP FE is now mostly API-driven for auth, dashboards, role scopes, integrations, notifications, analytics, AI insights, tasks, teams, employees, leave, attendance and organizations. The latest backend changelog introduced analytics dashboard, summary/burnout endpoints, integration sync policy updates, AI runtime fields, project-scoped AI insights, admin notifications and task.externalDeleted support.",
        fill=GOOD,
    )
    add_bullets(doc, [
        "FE changes from the latest backend update have been implemented and pushed: analytics dashboard UI, summary/burnout checker, integration sync controls by role, AI project insight scope, structured fullAnalysis parsing, admin notification endpoint, and task externalDeleted badge.",
        "Local verification passed: npm.cmd run lint, npm.cmd run build, and route smoke tests for /, /login, /dashboard, /analytics, /ai-insights, /integrations, /notifications, /tasks.",
        "Remaining work is mainly product hardening: cleaner role dashboards, stronger form validation against Swagger DTOs, full integration/Jira/GitHub testing with real credentials, and backend fixes for endpoints that still return server-side errors.",
    ])

    doc.add_heading("2. System Overview", level=1)
    add_table(doc, ["Layer", "Current Design", "Primary Responsibility"], [
        ["Frontend", "React + Vite product frontend", "Role-based SaaS UI, protected routes, API service layer, dark mode, notification UI, landing page, forms and CRUD screens."],
        ["Backend", "Spring Boot 3 / Java 21 / PostgreSQL / JWT", "Authentication, RBAC, organization/team/employee/task domain, integrations, analytics, notifications and dashboards."],
        ["AI Service / Runtime", "Backend AI runtime currently configured for Gemini according to changelog", "Generate insights, summarize risks, parse structured AI output, support RAG context."],
        ["Integrations", "GitHub, Jira, Internal provider enum", "Sync external issues/activity into operational records and tasks; store sync logs and counters."],
        ["Deployment", "FE on Vercel, BE on Render", "SPA routing requires Vercel rewrite; backend may cold-start and must return clear API errors."],
    ], widths=[1.2, 2.1, 3.2])

    doc.add_heading("3. Role-Based Product Scope", level=1)
    add_table(doc, ["Role", "Dashboard & Main Navigation", "Allowed Product Scope"], [
        ["Admin", "Dashboard, Organizations, Users, Teams, Sprints, Analytics, AI Insights, Integrations, Notifications, Settings.", "Platform control: manage organizations, users, teams, sprints, integrations, analytics, AI, notifications and security settings."],
        ["Manager", "Dashboard, Team Tasks, Employees, Team Analytics, AI Insights, Event Timeline, Attendance, Reports, Notifications, Settings.", "Team operations: managed teams, tasks, workload, attendance, reports and insight review. Can trigger integration sync only for owned team/project scope."],
        ["HR / People Ops", "Dashboard, Employees, Attendance, Leave Requests, Recruitment, People Analytics, AI Insights, Reports, Notifications, Settings.", "People operations: employees, attendance, leave, recruitment/onboarding, people analytics and people risk signals."],
        ["Employee", "Dashboard, My Tasks, Attendance, Leave Requests, Personal Analytics, Event Timeline, AI Insights, Notifications, Profile.", "Personal workspace: own tasks, attendance, leave, profile, personal analytics and personal AI insights. No integration sync trigger."],
    ], widths=[1.1, 2.4, 3.0])

    doc.add_heading("4. Latest Backend Changes Already Reflected in FE", level=1)
    add_table(doc, ["Backend Change", "FE Status", "Files Updated"], [
        ["GET /api/v1/analytics/dashboard", "Implemented in service and rendered in Analytics page with task totals, burnout risk, recent activity and AI insight summary.", "analyticsService.js, AnalyticsPage.jsx"],
        ["Analytics summary + burnout endpoints", "Service methods added; Analytics page has scope selector for my/team/employee summary and burnout checks.", "analyticsService.js, AnalyticsPage.jsx"],
        ["Integration sync status RUNNING/SUCCESS/FAILED and counters", "Sync logs table now displays errorMessage, totalFetched, totalCreated, totalUpdated and finishedAt.", "integrationService.js, IntegrationsPage.jsx"],
        ["POST /api/v1/integrations/sync", "Admin-only Sync All button added. Employee sync trigger hidden.", "integrationService.js, IntegrationsPage.jsx"],
        ["AI runtime status apiKeyConfigured", "AI Insights page displays provider/model and warning if API key is not configured.", "AIInsightsPage.jsx"],
        ["GET /api/v1/ai/insights/project/{projectId}", "Project ID scope added to AI Insights page.", "aiInsightService.js, AIInsightsPage.jsx"],
        ["Structured fullAnalysis JSON", "FE parses riskLevel, summary, reasons and recommendations instead of assuming old fields.", "AIInsightsPage.jsx"],
        ["GET /api/v1/notifications/admin/all", "Admin notification page uses admin endpoint; other roles use personal notifications.", "notificationService.js, NotificationPage.jsx"],
        ["Task externalDeleted", "Task list supports External deleted badge.", "TaskPage.jsx"],
    ], widths=[2.0, 3.0, 1.5])

    doc.add_heading("5. FE Fixes Still Recommended", level=1)
    add_table(doc, ["Priority", "Area", "Issue", "Recommended Fix"], [
        ["High", "Forms / DTOs", "Some create/update forms still expose raw UUID inputs and can feel technical.", "Replace raw UUID inputs with searchable selectors loaded from API. Keep UUID fields only in developer/admin advanced mode."],
        ["High", "Role UX", "Manager/HR/Employee dashboards may show empty states when backend has no matching scoped data or context.", "Load account context from /auth/me and profile endpoints consistently. Display role-specific empty states and hide actions not allowed by backend."],
        ["High", "Errors", "Some backend errors are technically correct but too raw for normal users.", "Keep technical details collapsible. Main message should be human: permission denied, missing context, backend validation failed, dependency conflict, or endpoint unavailable."],
        ["Medium", "Analytics", "Analytics dashboard now renders the new endpoint, but workloadByEmployee could be more useful.", "Add workload table/bar visual for workloadByEmployee; add filters for risk/team/employee."],
        ["Medium", "AI Insights", "Project-scoped insight needs project selector rather than pasted UUID.", "Add Project dropdown once project API returns list/search cleanly."],
        ["Medium", "Integrations", "GitHub/Jira sync testing requires real credentials/tokens and configured integration records.", "Add setup checklist UI and validation messages for provider, credentials, team/project ownership and sync log result."],
        ["Medium", "Internationalization", "User requested English/Vietnamese toggle earlier.", "Centralize UI strings in i18n resource files. Do not hardcode mixed language in component JSX."],
        ["Low", "Bundle", "Build warns chunk is over 500KB.", "Lazy-load landing/app heavy pages or route chunks after core functionality stabilizes."],
    ], widths=[0.75, 1.15, 2.3, 2.3])

    doc.add_heading("6. BE Fixes / Backend Notes", level=1)
    add_callout(
        doc,
        "Important",
        "These backend items are based on observed UI/API behavior and the new changelog. They should be verified directly in Swagger/Postman after the latest redeploy.",
        fill=WARN,
    )
    add_table(doc, ["Priority", "Backend Area", "Observed / Expected Problem", "Recommended Backend Fix"], [
        ["High", "Tasks API", "GET /api/v1/tasks returned server error: Hibernate ByteBuddy proxy type definition error.", "Return DTOs only and avoid exposing lazy JPA entities/proxies. Add @JsonIgnore or DTO mapping for nested relations."],
        ["High", "Integrations / Projects", "GET /api/v1/integrations and GET /api/v1/projects previously returned method GET not supported.", "If list pages exist in FE, expose supported GET list endpoints or update Swagger/FE to only use supported team/project-scoped endpoints."],
        ["High", "Organization delete", "Deleting organization failed due FK constraint team_organization_id_fkey.", "Return 409 Conflict with business message or implement cascade/soft delete policy. FE should not show this as generic 500."],
        ["High", "RBAC consistency", "Manager/Employee often fail when calling scoped endpoints if backend account context is incomplete.", "Ensure /auth/me or profile response includes userId, employeeId, role, organizationId, teamId/managedTeams where applicable."],
        ["Medium", "Integration Sync", "Sync policy is role-sensitive and external credential-sensitive.", "Return clear validation errors for missing token, bad Jira/GitHub credentials, no team/project ownership, or sync already running."],
        ["Medium", "AI Runtime", "If apiKeyConfigured=false, generate endpoints may fail.", "Return deterministic 400/503 with message, provider/model, and whether API key/RAG config is missing."],
        ["Medium", "Response shape", "FE supports ApiResponse wrapper, arrays, data/result/payload/content/items.", "Keep ApiResponse consistent: success, message, data. Avoid returning raw exceptions as data."],
        ["Low", "Pagination", "Large list pages need stable pagination response.", "Standardize page shape: content, totalElements, totalPages, number, size."],
    ], widths=[0.75, 1.3, 2.25, 2.2])

    doc.add_heading("7. API Coverage Snapshot", level=1)
    add_table(doc, ["Domain", "FE Coverage", "Next Check"], [
        ["Auth", "Login/register/logout/me token flow exists; protected routes use JWT.", "Verify backend login errors return credential-specific messages and role/account context."],
        ["Organizations", "Admin CRUD UI exists and calls API.", "Handle FK delete conflicts with friendly 409 copy."],
        ["Teams", "Team list/create/update/delete/assign member UI exists.", "Replace UUID inputs with selectors and respect manager/admin/HR permissions."],
        ["Employees", "Directory/profile update UI exists.", "Ensure employee role does not see global directory."],
        ["Tasks", "Task table/forms/status/delete use real service and safe field mapping.", "Backend task list serialization issue must be fixed."],
        ["Attendance", "Role-based history/overview/check-in/out service layer exists.", "Test with accounts that have employeeId/teamId/organizationId."],
        ["Leave", "Create/approve/reject wired by role.", "Confirm HR/Manager permission matrix and request body names."],
        ["Notifications", "Personal and admin all endpoints covered.", "Verify owner guard messages for mark/delete errors."],
        ["Analytics/Burnout", "Dashboard, workload history, summary and burnout endpoints covered.", "Add richer workloadByEmployee UI and team/employee selectors."],
        ["AI Insights/Suggestions", "Runtime status, insights, project scope, generate and adopt suggestion service covered.", "Test generate/adopt with real Gemini config and employee/project IDs."],
        ["Integrations", "Provider config, OAuth links, sync, sync logs and sync-all covered.", "Real GitHub/Jira credentials required to validate task/commit sync end-to-end."],
    ], widths=[1.2, 2.8, 2.5])

    doc.add_heading("8. Test Plan After Next Deploy", level=1)
    add_numbered(doc, [
        "Login with Admin, Manager, HR and Employee accounts. Confirm role is derived from the backend account, not manually switched for normal users.",
        "Admin: open Dashboard, Organizations, Teams, Sprints, Analytics, AI Insights, Integrations and Notifications. Validate data loads or clean empty/error states.",
        "Manager: verify managed team tasks, managed team attendance, managed leaves, analytics summary/team burnout and integration sync visibility.",
        "HR: verify employee directory, attendance overview, leave review, recruitment placeholder, people analytics and notifications.",
        "Employee: verify My Tasks, My Attendance, My Leave Requests, My AI Insights, Notifications and Profile only.",
        "Integration: configure GitHub/Jira provider with real credentials, run sync, inspect sync logs for RUNNING/SUCCESS/FAILED and fetched/created/updated counters.",
        "AI: check runtime status, generate employee insight with valid employeeId, inspect fullAnalysis reasons/recommendations, test project insight endpoint.",
        "Regression: run npm.cmd run lint, npm.cmd run build, reload deployed Vercel protected routes, and verify SPA rewrite works.",
    ])

    doc.add_heading("9. Recommended Next Sprint Order", level=1)
    add_table(doc, ["Order", "Work Item", "Owner", "Definition of Done"], [
        ["1", "Fix backend task serialization and unsupported GET list endpoints.", "Backend", "Task and project/integration list pages load without 500 errors."],
        ["2", "Finalize account context contract.", "Backend + FE", "/auth/me or profile returns role, employeeId, organizationId, team/team ownership."],
        ["3", "Replace raw UUID form fields with API selectors.", "Frontend", "Team/task/employee/org/project forms are usable by normal users."],
        ["4", "Run role-by-role QA with real accounts.", "FE + QA", "Each role sees only its expected sidebar, actions and API scope."],
        ["5", "Run GitHub/Jira sync test with credentials.", "Backend + FE", "External issues/commits sync, local tasks update, sync logs show counters."],
        ["6", "Run AI generate/adopt suggestion test.", "Backend + FE", "AI runtime status valid, insight generated, fullAnalysis parsed, suggestion adoption persists."],
    ], widths=[0.7, 2.3, 1.1, 2.4])

    doc.add_heading("10. Notes for Product Readiness", level=1)
    add_bullets(doc, [
        "Avoid showing raw endpoint paths in user-facing dashboards except in admin/developer diagnostics.",
        "Use backend messages, but translate them into human-readable UI copy; keep raw technical details collapsible.",
        "Never fake successful create/update/delete actions. If backend rejects a request, keep the form open and show the backend reason.",
        "Keep role selection locked to backend account role in production. Admin-only test override can exist behind a clear developer/test flag.",
        "For Vercel, keep SPA rewrite in vercel.json so direct reload on /login, /dashboard and other routes does not return 404.",
    ])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("FOREP - System overview and FE/BE action plan")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string(MUTED)

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
